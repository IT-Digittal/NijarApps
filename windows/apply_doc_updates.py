"""Aplica las actualizaciones acordadas a los dos Word de SAT:
- ChecklistSmokeTestDTINijar.docx
- EstadoyPlandePruebasDTINijar.docx

Sobrescribe los .docx originales tras hacer copia de seguridad .bak.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

REPO = Path(__file__).resolve().parent.parent


def set_para_text(paragraph, text: str) -> None:
    """Reemplaza el texto del párrafo preservando el formato del primer run."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    """Inserta un nuevo párrafo justo después de ``paragraph``."""
    new_p = deepcopy(paragraph._element)
    # Limpia los hijos del nuevo párrafo (mantiene pPr para conservar estilo)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = paragraph.part.document.styles[style]
    p.add_run(text)
    return p


# ============================================================================
# 1) Checklist
# ============================================================================

def update_checklist() -> None:
    path = REPO / "ChecklistSmokeTestDTINijar.docx"
    shutil.copy2(path, path.with_suffix(".docx.bak"))
    d = Document(path)

    for i, p in enumerate(d.paragraphs):
        txt = p.text

        # Línea 7 — Login: añadir nota de formatos de API
        if txt.startswith("Login panel:"):
            set_para_text(
                p,
                "Login panel: admin@nijar.es / CambiarEnPrimerArranque#2026  "
                "· API JSON: POST /auth/login {\"email\":\"…\",\"password\":\"…\"}  "
                "· API OAuth2 form: POST /auth/token (username=email, password)",
            )

        # Línea 16 — MAPE puede ser null
        elif txt.startswith("GET /prediccion/validacion?metrica=chatbot"):
            set_para_text(
                p,
                "GET /prediccion/validacion?metrica=chatbot → MAPE "
                "(puede ser null si la serie tiene muchos valles a cero — ver `nota`/`n_evaluable` en la respuesta)",
            )

        # Línea 22 — Incidencias rango relativo
        elif txt.startswith("GET /incidencias?desde="):
            set_para_text(
                p,
                "GET /incidencias?desde=<inicio_mes_anterior>T00:00:00Z&hasta=<inicio_mes_actual>T00:00:00Z → 8 incidencias "
                "(el seed las crea relativas al mes natural anterior a la fecha de arranque)",
            )

        # Línea 23 — ANS añadir valor esperado
        elif txt.startswith("GET /incidencias/ans?desde="):
            set_para_text(
                p,
                "GET /incidencias/ans?desde=…&hasta=… → cumplimiento por severidad "
                "(esperado: alta ≈ 66,7 % — 3 altas, 2 cumplen)",
            )

        # Línea 24 — Monthly report URL correcta
        elif txt.startswith("GET /dashboards/monthly-report"):
            set_para_text(
                p,
                "GET /dashboards/reports/monthly?year=<aaaa>&month=<mm> → disponibilidad real",
            )

    # Añadir nota nueva al bloque de Notas/incidencias detectadas y otra
    # bajo "Backend" sobre auto-seed. La forma robusta es localizar el heading
    # "Notas / incidencias detectadas" y añadir párrafo justo debajo.
    for p in d.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("Notas / incidencias"):
            insert_paragraph_after(
                p,
                "Los seeds (recursos, sensores, FAQs, eventos, observaciones, "
                "incidencias del mes anterior y contexto histórico) se cargan "
                "automáticamente al arrancar la API (flag RUN_SEEDS_ON_STARTUP). "
                "Para forzar reload: `docker compose exec api python -m nijar_dti.data.seed_loader`.",
                style="Compact",
            )
            break

    d.save(path)
    print(f"✓ Checklist actualizada ({path.name})")


# ============================================================================
# 2) Estado y plan de pruebas
# ============================================================================

def update_estado() -> None:
    path = REPO / "EstadoyPlandePruebasDTINijar.docx"
    shutil.copy2(path, path.with_suffix(".docx.bak"))
    d = Document(path)

    # ---- Párrafos ----
    auto_seed_added = False
    for p in d.paragraphs:
        txt = p.text

        # Datos demo: añadir contexto histórico
        if txt.startswith("Datos demo cargados al arrancar:"):
            set_para_text(
                p,
                "Datos demo cargados al arrancar: 14 recursos, 9 sensores, 105 FAQs, "
                "eventos/observaciones/opiniones/visitas/chatbot e incidencias del mes "
                "anterior, y contexto histórico (INE EOH/Frontur/Egatur, Junta, AENA) "
                "en dry-run.",
            )

        # Nota sobre seed manual → ahora automático
        elif txt.startswith("Si una sección del panel sale vacía"):
            set_para_text(
                p,
                "Los seeds se cargan automáticamente al arrancar la API (busca en "
                "logs: «Ejecutando seed_loader idempotente al arranque» → «seed_loader "
                "completado»). Para forzar reload: "
                "docker compose exec api python -m nijar_dti.data.seed_loader. "
                "El loader es idempotente y refresca los eventos demo cuando sus fechas "
                "han caducado.",
            )

        # curl de obtención de token
        elif "curl -s -X POST localhost:8000/api/v1/auth/login" in txt and "username=admin" in txt:
            set_para_text(
                p,
                "TOKEN=$(curl -s -X POST localhost:8000/api/v1/auth/token \\\n"
                "  -d 'username=admin@nijar.es&password=CambiarEnPrimerArranque#2026' "
                "| jq -r .access_token)\n"
                "# Alternativa con JSON: POST /auth/login -H 'Content-Type: application/json' "
                "-d '{\"email\":\"…\",\"password\":\"…\"}'",
            )

        # Heading "Backfill de contexto" y bloque siguiente
        elif txt.strip().startswith("Backfill de contexto"):
            set_para_text(
                p,
                "Backfill de contexto (ahora automático en el arranque)",
            )

        elif "python -m nijar_dti.workers.contexto_backfill" in txt and "POST /data/contexto/ingest" in txt:
            set_para_text(
                p,
                "# Ya no es necesario en arranque normal: el seed_loader lo ejecuta.\n"
                "# Para forzar regeneración tras vaciar `contexto_turistico`:\n"
                "docker compose exec api python -m nijar_dti.workers.contexto_backfill --dry-run --output /tmp/ctx.json\n"
                "# y luego POST /data/contexto/ingest con ese JSON (rol administrador_tic o analista_datos).",
            )

        # P09/P10 nota — añadir aviso de MAPE null
        elif txt.startswith("En P09 y P10, con los datos demo"):
            set_para_text(
                p,
                "En P09 y P10, con los datos demo: disponibilidad media ≈ 99,3 %, ANS "
                "de \"alta\" ≈ 66,7 % (1 incumplimiento), y la predicción debe dibujar "
                "la curva + bandas. Nota: el MAPE de validación puede devolver null "
                "cuando la serie demo tiene muchos valles a cero — está documentado en "
                "la propia respuesta (campos `nota` y `n_evaluable`).",
            )

        # Avisos de lint: añadir nota sobre cwd
        elif txt.startswith("Avisos de lint esperados"):
            set_para_text(
                p,
                "Avisos de lint esperados (no son errores): UP017 (timezone.utc), N818 "
                "(NotFound/Conflict), S311 (random en datos demo), E501 en data/seeds/* "
                "(exento). No bloquean. Nota: `tests/test_coherencia_paquete.py` valida "
                "presencia de `frontend/`, `infra/`, `.github/`, `rasa/`, `docs/` y "
                "`README.md`; lánzalo desde la raíz del repo (no desde `/app` dentro del "
                "contenedor API, que es read-only).",
            )

    # Añadir bullet nuevo en Sección 1 sobre auto-seed (tras el "Datos demo cargados…")
    for p in d.paragraphs:
        if p.text.startswith("Datos demo cargados al arrancar:") and not auto_seed_added:
            insert_paragraph_after(
                p,
                "Seeds idempotentes en lifespan de la API: al arrancar, FastAPI llama "
                "a seed_loader.run() y refresca eventos demo cuyas fechas hayan "
                "caducado (mantiene IDs, actualiza fecha/i18n). Configurable con "
                "RUN_SEEDS_ON_STARTUP (default True en dev).",
                style="Compact",
            )
            auto_seed_added = True
            break

    # Añadir referencia nueva en Sección 8
    for p in d.paragraphs:
        if p.text.startswith("Índice general (dossier)"):
            insert_paragraph_after(
                p,
                "Configuración auto-seed: src/nijar_dti/config.py (flag run_seeds_on_startup).",
                style="Compact",
            )
            break

    # ---- Tablas ----
    backend_tbl = d.tables[2]
    for row in backend_tbl.rows:
        cells_text = [c.text for c in row.cells]
        if not cells_text:
            continue
        rid = cells_text[0].strip()

        if rid == "B01":
            for p in row.cells[2].paragraphs:
                if "POST /auth/login" in p.text:
                    set_para_text(
                        p,
                        "POST /auth/login (JSON: email/password)  ·  "
                        "POST /auth/token (form: username/password)  ·  GET /auth/me",
                    )
                    break

        elif rid == "B14":
            for p in row.cells[3].paragraphs:
                if p.text.startswith("MAPE"):
                    set_para_text(
                        p,
                        "MAPE + cumple_umbral (MAPE puede ser null si la serie tiene "
                        "muchos ceros — ver `nota` y `n_evaluable`)",
                    )
                    break

        elif rid == "B12":
            for p in row.cells[3].paragraphs:
                if "Serie" in p.text:
                    set_para_text(
                        p,
                        "Serie histórica (backfill automático al arrancar; campo de "
                        "respuesta `puntos`)",
                    )
                    break

        elif rid == "B21":
            for p in row.cells[2].paragraphs:
                if "GET /incidencias?desde=" in p.text:
                    set_para_text(
                        p,
                        "GET /incidencias?desde=<inicio_mes_anterior>T00:00:00Z"
                        "&hasta=<inicio_mes_actual>T00:00:00Z",
                    )
                    break

        elif rid == "B23":
            for p in row.cells[2].paragraphs:
                if "monthly-report" in p.text:
                    set_para_text(
                        p,
                        "GET /dashboards/reports/monthly?year=<aaaa>&month=<mm>",
                    )
                    break

    d.save(path)
    print(f"✓ Estado actualizado ({path.name})")


if __name__ == "__main__":
    update_checklist()
    update_estado()
